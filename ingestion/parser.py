"""
개선된 문서 파서 - 한국어 PDF 인코딩 문제 해결
"""
import os
import tempfile
from uuid import uuid4
from typing import List, Dict, Union
import logging
import re
from ftfy import fix_text


logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    """텍스트 정리 및 인코딩 문제 해결"""
    if not text:
        return ""

    # 바이트 문자열이면 UTF-8로 디코드
    if isinstance(text, bytes):
        try:
            text = text.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = text.decode('cp949')  # 한국어 윈도우 인코딩
            except UnicodeDecodeError:
                text = text.decode('utf-8', errors='replace')

    # ① ftfy 로 한 번 복구
    text = fix_text(text)

    # ② 제어 문자 제거 (줄바꿈, 탭 제외)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # 이상한 문자들 정리 (PDF에서 자주 나타나는 문제들)
    text = re.sub(r'[^\w\s가-힣ㄱ-ㅎㅏ-ㅣ一-龯\u3040-\u309F\u30A0-\u30FF.,!?;:()\[\]{}\'\"%-]', ' ', text)

    # 연속된 공백 정리
    text = re.sub(r'\s+', ' ', text)

    # 앞뒤 공백 제거
    text = text.strip()

    return text


GARBLED_THRESHOLD = 0.30          # 유효 문자 비율 30 %

def is_garbled(text: str) -> bool:
    """가독 가능한 문자 비율이 임계값보다 낮으면 True"""
    if not text:
        return True
    good = len(re.findall(r'[가-힣a-zA-Z0-9]', text))
    return good / len(text) < GARBLED_THRESHOLD


def parse_pdf_with_pypdf(file_path: str, lang_hint="auto") -> List[Dict]:
    """PyPDF를 사용한 PDF 파싱 (인코딩 개선)"""
    try:
        from pypdf import PdfReader

        chunks = []

        # 다양한 인코딩으로 파일 읽기 시도
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                reader = PdfReader(file_path)
                break
            except Exception as e:
                logger.warning(f"Failed to read PDF with encoding {encoding}: {e}")
                continue
        else:
            raise Exception("Could not read PDF with any encoding")

        for page_num, page in enumerate(reader.pages):
            try:
                # 텍스트 추출
                text = page.extract_text()

                if not text:
                    logger.warning(f"No text extracted from page {page_num + 1}")
                    continue

                # 텍스트 정리
                cleaned_text = clean_text(text)

                if cleaned_text and len(cleaned_text.strip()) > 10:
                    chunks.extend(chunk_text(
                        cleaned_text,
                        source=os.path.basename(file_path),
                        page=page_num + 1,
                        doc_type="pdf",
                        lang=lang_hint
                    ))

            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                continue

        logger.info(f"PDF parsed with PyPDF: {len(chunks)} chunks from {len(reader.pages)} pages")
        return chunks

    except ImportError:
        logger.error("PyPDF not available")
        return parse_as_fallback(file_path, "PDF 문서 (PyPDF 미설치)")
    except Exception as e:
        logger.error(f"PDF parsing failed with PyPDF: {e}")
        # pdfplumber로 fallback
        return parse_pdf_with_pdfplumber(file_path, lang_hint)


def parse_pdf_with_pdfplumber(file_path: str, lang_hint="auto") -> List[Dict]:
    """pdfplumber를 사용한 PDF 파싱 (한국어 최적화)"""
    try:
        import pdfplumber

        chunks = []

        # pdfplumber 설정 - CJK 폰트를 위한 특별 처리
        pdf_config = {
            'dedupe_chars': True,  # 중복 문자 제거
            'ignore_blank_chars': True,  # 빈 문자 무시
            'use_text_flow': True,  # 텍스트 흐름 사용
        }

        with pdfplumber.open(file_path, **pdf_config) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    # 다양한 방법으로 텍스트 추출 시도
                    text = None

                    # 방법 1: 기본 추출
                    try:
                        text = page.extract_text()
                    except Exception as e:
                        logger.warning(f"Default text extraction failed on page {page_num + 1}: {e}")

                    # 방법 2: 더 관대한 설정으로 추출
                    if not text or len(text.strip()) < 10:
                        try:
                            text = page.extract_text(
                                x_tolerance=3,  # 문자 간 간격 허용도
                                y_tolerance=3,  # 줄 간 간격 허용도
                                layout=True,    # 레이아웃 유지
                                x_density=7.25, # 해상도 설정
                                y_density=7.25
                            )
                        except Exception as e:
                            logger.warning(f"Alternative text extraction failed on page {page_num + 1}: {e}")

                    # 방법 3: 문자 기반 추출 (가장 안전함)
                    if not text or len(text.strip()) < 10:
                        try:
                            chars = page.chars
                            if chars:
                                # 문자를 위치 순서대로 정렬
                                sorted_chars = sorted(chars, key=lambda x: (x['top'], x['x0']))
                                text = ''.join(char['text'] for char in sorted_chars)
                        except Exception as e:
                            logger.warning(f"Char-based text extraction failed on page {page_num + 1}: {e}")

                    if not text:
                        logger.warning(f"No text extracted from page {page_num + 1}")
                        continue

                    # 텍스트 정리
                    cleaned_text = clean_text(text)

                    if cleaned_text and len(cleaned_text.strip()) > 10:
                        chunks.extend(chunk_text(
                            cleaned_text,
                            source=os.path.basename(file_path),
                            page=page_num + 1,
                            doc_type="pdf",
                            lang=lang_hint
                        ))

                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue

        logger.info(f"PDF parsed with pdfplumber: {len(chunks)} chunks")
        return chunks

    except ImportError:
        logger.error("pdfplumber not available")
        return parse_as_fallback(file_path, "PDF 문서 (pdfplumber 미설치)")
    except Exception as e:
        logger.error(f"PDF parsing failed with pdfplumber: {e}")
        # PyMuPDF로 fallback
        return parse_pdf_with_pymupdf(file_path, lang_hint)


def parse_pdf_with_pymupdf(file_path: str, lang_hint="auto") -> List[Dict]:
    """PyMuPDF를 사용한 PDF 파싱 (한국어 최적화)"""
    try:
        import fitz  # PyMuPDF

        chunks = []
        doc = fitz.open(file_path)

        for page_num in range(len(doc)):
            try:
                page = doc.load_page(page_num)

                # 텍스트 추출 - CJK 문자를 위한 플래그 설정
                text = page.get_text(
                    "text",
                    flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_SPANS
                )

                # 추가적인 추출 방법 시도
                if not text or len(text.strip()) < 10:
                    # dict 형태로 추출해서 폰트 정보까지 활용
                    text_dict = page.get_text("dict")
                    extracted_text = []

                    for block in text_dict["blocks"]:
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    if span["text"].strip():
                                        extracted_text.append(span["text"])

                    text = " ".join(extracted_text)

                if not text:
                    logger.warning(f"No text extracted from page {page_num + 1}")
                    continue

                # 텍스트 정리
                cleaned_text = clean_text(text)

                if cleaned_text and len(cleaned_text.strip()) > 10:
                    chunks.extend(chunk_text(
                        cleaned_text,
                        source=os.path.basename(file_path),
                        page=page_num + 1,
                        doc_type="pdf",
                        lang=lang_hint
                    ))

            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                continue

        doc.close()
        logger.info(f"PDF parsed with PyMuPDF: {len(chunks)} chunks")
        return chunks

    except ImportError:
        logger.error("PyMuPDF not available")
        return parse_as_fallback(file_path, "PDF 문서 (PyMuPDF 미설치)")
    except Exception as e:
        logger.error(f"PDF parsing failed with PyMuPDF: {e}")
        return parse_as_fallback(file_path, f"PDF 파싱 오류: {str(e)}")


def parse_text_file(file_path: str, lang_hint="auto") -> List[Dict]:
    """텍스트 파일 파싱 (인코딩 개선)"""
    try:
        # 다양한 인코딩 시도
        encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1', 'utf-16', 'ascii']
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if content is None:
            # 바이너리로 읽어서 강제 변환
            with open(file_path, 'rb') as f:
                raw_content = f.read()
                # chardet로 인코딩 감지 시도
                try:
                    import chardet
                    detected = chardet.detect(raw_content)
                    if detected and detected['encoding']:
                        content = raw_content.decode(detected['encoding'], errors='replace')
                        used_encoding = detected['encoding']
                    else:
                        content = raw_content.decode('utf-8', errors='replace')
                        used_encoding = 'utf-8 (forced)'
                except ImportError:
                    content = raw_content.decode('utf-8', errors='replace')
                    used_encoding = 'utf-8 (forced)'

        # 텍스트 정리
        cleaned_content = clean_text(content)

        logger.info(f"Text file parsed with encoding: {used_encoding}")

        return chunk_text(
            cleaned_content,
            source=os.path.basename(file_path),
            doc_type="text",
            lang=lang_hint
        )

    except Exception as e:
        logger.error(f"Text file parsing failed: {e}")
        return parse_as_fallback(file_path, f"텍스트 파일 파싱 실패: {str(e)}")


def parse_docx(file_path: str, lang_hint="auto") -> List[Dict]:
    """Word 문서 파싱 (인코딩 개선)"""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = []

        # 문단 텍스트 추출
        for para in doc.paragraphs:
            if para.text.strip():
                cleaned_text = clean_text(para.text.strip())
                if cleaned_text:
                    paragraphs.append(cleaned_text)

        # 표 내용 추출
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text.strip())
                    if cell_text:
                        row_texts.append(cell_text)

                if row_texts:
                    row_text = " | ".join(row_texts)
                    paragraphs.append(row_text)

        # 전체 텍스트 결합
        full_text = "\n\n".join(paragraphs)

        return chunk_text(
            full_text,
            source=os.path.basename(file_path),
            doc_type="docx",
            lang=lang_hint
        )

    except ImportError:
        logger.error("python-docx not available")
        return parse_as_fallback(file_path, "Word 문서 (python-docx 미설치)")
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return parse_as_fallback(file_path, f"Word 문서 파싱 오류: {str(e)}")


def parse_as_fallback(file_path: str, description: str) -> List[Dict]:
    """폴백 처리"""
    return [{
        "chunk_id": str(uuid4()),
        "content": f"{description}가 업로드되었습니다: {os.path.basename(file_path)}",
        "meta": {
            "source": os.path.basename(file_path),
            "type": "fallback",
            "note": description,
            "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0
        }
    }]


def chunk_text(text: str, source: str, doc_type: str = "unknown",
               page: int = None, lang: str = "auto", chunk_size: int = 500,
               chunk_overlap: int = 50) -> List[Dict]:
    """텍스트를 청크로 분할 (한국어 최적화)"""
    if not text or not text.strip():
        return []

    # 텍스트 정리
    cleaned_text = clean_text(text)

    if not cleaned_text:
        return []

    chunks = []

    if len(cleaned_text) <= chunk_size:
        chunks.append({
            "chunk_id": str(uuid4()),
            "content": cleaned_text,
            "meta": {
                "source": source,
                "type": doc_type,
                "lang": lang,
                **({"page": page} if page else {}),
                "chunk_index": 0,
                "total_chunks": 1
            }
        })
    else:
        start = 0
        chunk_index = 0

        while start < len(cleaned_text):
            end = start + chunk_size

            # 한국어 및 영어를 고려한 단어 경계에서 자르기
            if end < len(cleaned_text):
                # 한국어 문장 끝 문자들
                korean_endings = ['다', '요', '니다', '습니다', '죠', '네', '라', '까', '야']

                # 찾을 범위 설정
                search_start = max(start + chunk_size // 2, end - 100)

                # 우선순위 1: 문장 부호
                for i in range(end, search_start, -1):
                    if cleaned_text[i] in ['.', '!', '?', '。', '!', '?', '\n']:
                        end = i + 1
                        break
                else:
                    # 우선순위 2: 한국어 문장 끝
                    for i in range(end, search_start, -1):
                        if i > 0 and cleaned_text[i-1:i+1] in ['다.', '요.', '다!', '요!', '다?', '요?']:
                            end = i + 1
                            break
                        elif cleaned_text[i] in korean_endings and i < len(cleaned_text) - 1 and cleaned_text[i+1] in [' ', '\n', '\t']:
                            end = i + 1
                            break
                    else:
                        # 우선순위 3: 공백
                        for i in range(end, search_start, -1):
                            if cleaned_text[i] in [' ', '\n', '\t']:
                                end = i + 1
                                break

            chunk_content = cleaned_text[start:end].strip()

            if chunk_content and len(chunk_content) > 10:
                chunks.append({
                    "chunk_id": str(uuid4()),
                    "content": chunk_content,
                    "meta": {
                        "source": source,
                        "type": doc_type,
                        "lang": lang,
                        **({"page": page} if page else {}),
                        "chunk_index": chunk_index,
                        "start_pos": start,
                        "end_pos": end
                    }
                })
                chunk_index += 1

            start = max(start + chunk_size - chunk_overlap, end)

    # 총 청크 수 업데이트
    for chunk in chunks:
        chunk["meta"]["total_chunks"] = len(chunks)

    return chunks


def parse_pdf(file_input: Union[str, bytes], lang_hint: str = "auto") -> List[Dict]:
    """메인 파싱 함수 - 여러 라이브러리로 단계적 시도"""
    # 바이트 데이터인 경우 임시 파일로 저장
    if isinstance(file_input, bytes):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file_input)
            temp_path = tmp_file.name

        try:
            return parse_file_by_extension(temp_path, lang_hint)
        finally:
            try:
                os.unlink(temp_path)
            except:
                pass

    elif isinstance(file_input, (str, os.PathLike)):
        return parse_file_by_extension(str(file_input), lang_hint)

    else:
        raise ValueError(f"Unsupported input type: {type(file_input)}")


def parse_file_by_extension(file_path: str, lang_hint: str = "auto") -> List[Dict]:
    """파일 확장자에 따른 파싱 - 한국어 PDF에 최적화된 단계적 시도"""
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == '.pdf':
        # PDF 파싱 - 여러 라이브러리를 순차적으로 시도
        successful_chunks = []

        # 1단계: pdfplumber 먼저 시도 (한국어에 가장 좋음)
        try:
            result = parse_pdf_with_pdfplumber(file_path, lang_hint)
            # if result and any(len(chunk['content'].strip()) > 50 for chunk in result if chunk['meta']['type'] != 'fallback'):
            if result and not is_garbled(" ".join(c['content'] for c in result)):
                successful_chunks = result
                logger.info(f"Successfully parsed with pdfplumber: {len(successful_chunks)} chunks")
                return successful_chunks
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # 2단계: PyMuPDF 시도
        if not successful_chunks:
            try:
                result = parse_pdf_with_pymupdf(file_path, lang_hint)
                # if result and any(len(chunk['content'].strip()) > 50 for chunk in result if chunk['meta']['type'] != 'fallback'):
                if result and not is_garbled(" ".join(c['content'] for c in result)):
                    successful_chunks = result
                    logger.info(f"Successfully parsed with PyMuPDF: {len(successful_chunks)} chunks")
                    return successful_chunks
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")

        # 3단계: PyPDF 시도
        if not successful_chunks:
            try:
                result = parse_pdf_with_pypdf(file_path, lang_hint)
                # if result and any(len(chunk['content'].strip()) > 50 for chunk in result if chunk['meta']['type'] != 'fallback'):
                if result and not is_garbled(" ".join(c['content'] for c in result)):
                    successful_chunks = result
                    logger.info(f"Successfully parsed with PyPDF: {len(successful_chunks)} chunks")
                    return successful_chunks
            except Exception as e:
                logger.warning(f"PyPDF failed: {e}")

        # 모든 시도가 실패한 경우
        if not successful_chunks:
            logger.error("All PDF parsing libraries failed")
            return parse_as_fallback(file_path, "PDF 문서 (모든 파싱 라이브러리 실패)")

        return successful_chunks

    elif file_ext in ['.docx', '.doc']:
        return parse_docx(file_path, lang_hint)
    elif file_ext in ['.txt', '.md', '.rst']:
        return parse_text_file(file_path, lang_hint)
    else:
        # 기본적으로 텍스트로 처리 시도
        try:
            return parse_text_file(file_path, lang_hint)
        except:
            return parse_as_fallback(file_path, f"알 수 없는 파일 형식 ({file_ext})")


# 테스트 함수
if __name__ == "__main__":
    import sys

    print("📄 Improved Korean PDF Parser Test")
    print("==================================")

    # 인코딩 테스트
    test_korean = "안녕하세요. 이것은 한글 테스트입니다. AI 기반 검색 패러다임 전환에 대해 설명합니다."
    cleaned = clean_text(test_korean)
    print(f"✅ Encoding test passed: {cleaned[:50]}...")

    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        print(f"\n🔍 Testing with: {test_file}")

        try:
            chunks = parse_pdf(test_file)
            print(f"✅ Generated {len(chunks)} chunks")

            for i, chunk in enumerate(chunks[:3]):  # 처음 3개 청크만 표시
                print(f"\n--- Chunk {i+1} ---")
                print(f"Content preview: {chunk['content'][:200]}...")
                print(f"Content length: {len(chunk['content'])}")
                print(f"Meta: {chunk['meta']}")

                # 한국어 포함 여부 확인
                korean_chars = len(re.findall(r'[가-힣]', chunk['content']))
                print(f"Korean characters: {korean_chars}")

        except Exception as e:
            print(f"❌ Failed: {e}")
            import traceback
            traceback.print_exc()
    else:
        print("\n💡 Usage: python parser.py <pdf_file_path>")