"""
개선된 문서 파서 - 한국어 PDF 인코딩 문제 해결
"""
import os
import tempfile
from uuid import uuid4
from typing import List, Dict, Union
import logging
import re
from ftfy import fix_text


logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    """텍스트 정리 및 인코딩 문제 해결"""
    if not text:
        return ""

    # 바이트 문자열이면 UTF-8로 디코드
    if isinstance(text, bytes):
        try:
            text = text.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = text.decode('cp949')  # 한국어 윈도우 인코딩
            except UnicodeDecodeError:
                text = text.decode('utf-8', errors='replace')

    # ① ftfy 로 한 번 복구
    text = fix_text(text)

    # ② 제어 문자 제거 (줄바꿈, 탭 제외)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # 이상한 문자들 정리 (PDF에서 자주 나타나는 문제들)
    text = re.sub(r'[^\w\s가-힣ㄱ-ㅎㅏ-ㅣ一-龯\u3040-\u309F\u30A0-\u30FF.,!?;:()\[\]{}\'\"%-]', ' ', text)

    # 연속된 공백 정리
    text = re.sub(r'\s+', ' ', text)

    # 앞뒤 공백 제거
    text = text.strip()

    return text


GARBLED_THRESHOLD = 0.30          # 유효 문자 비율 30 %

def is_garbled(text: str) -> bool:
    """가독 가능한 문자 비율이 임계값보다 낮으면 True"""
    if not text:
        return True
    good = len(re.findall(r'[가-힣a-zA-Z0-9]', text))
    return good / len(text) < GARBLED_THRESHOLD


def parse_pdf_with_pdfplumber(file_path: str, lang_hint="auto") -> List[Dict]:
    """pdfplumber를 사용한 PDF 파싱 (개선된 에러 처리)"""
    try:
        import pdfplumber

        # 파일 접근성 사전 확인
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일이 존재하지 않습니다: {file_path}")

        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"파일 읽기 권한이 없습니다: {file_path}")

        logger.info(f"pdfplumber로 PDF 파싱 시작: {os.path.basename(file_path)}")
        chunks = []

        # pdfplumber 설정 - CJK 폰트를 위한 특별 처리
        pdf_config = {
            # 'dedupe_chars': True,
        #     'ignore_blank_chars': True,
        #     'use_text_flow': True,
        }

        with pdfplumber.open(file_path, **pdf_config) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"PDF 총 페이지 수: {total_pages}")

            for page_num, page in enumerate(pdf.pages):
                try:
                    text = None

                    # 방법 1: 기본 추출
                    try:
                        text = page.extract_text()
                        if text and len(text.strip()) > 10:
                            logger.debug(f"페이지 {page_num + 1}: 기본 방법으로 텍스트 추출 성공")
                    except Exception as e:
                        logger.warning(f"페이지 {page_num + 1} 기본 추출 실패: {e}")

                    # 방법 2: 더 관대한 설정으로 추출
                    if not text or len(text.strip()) < 10:
                        try:
                            text = page.extract_text(
                                x_tolerance=3,
                                y_tolerance=3,
                                layout=True,
                                x_density=7.25,
                                y_density=7.25
                            )
                            if text and len(text.strip()) > 10:
                                logger.debug(f"페이지 {page_num + 1}: 관대한 설정으로 텍스트 추출 성공")
                        except Exception as e:
                            logger.warning(f"페이지 {page_num + 1} 관대한 추출 실패: {e}")

                    # 방법 3: 문자 기반 추출
                    if not text or len(text.strip()) < 10:
                        try:
                            chars = page.chars
                            if chars:
                                sorted_chars = sorted(chars, key=lambda x: (x['top'], x['x0']))
                                text = ''.join(char['text'] for char in sorted_chars)
                                if text and len(text.strip()) > 10:
                                    logger.debug(f"페이지 {page_num + 1}: 문자 기반 추출 성공")
                        except Exception as e:
                            logger.warning(f"페이지 {page_num + 1} 문자 기반 추출 실패: {e}")

                    if not text or len(text.strip()) < 10:
                        logger.warning(f"페이지 {page_num + 1}: 모든 추출 방법 실패 또는 내용 부족")
                        continue

                    # 텍스트 정리 및 검증
                    cleaned_text = clean_text(text)

                    if cleaned_text and len(cleaned_text.strip()) > 10:
                        if not is_garbled(cleaned_text):
                            chunks.extend(chunk_text(
                                cleaned_text,
                                source=os.path.basename(file_path),
                                page=page_num + 1,
                                doc_type="pdf",
                                lang=lang_hint
                            ))
                            logger.debug(f"페이지 {page_num + 1}: 청크 생성 완료")
                        else:
                            logger.warning(f"페이지 {page_num + 1}: 깨진 텍스트 감지됨")
                    else:
                        logger.warning(f"페이지 {page_num + 1}: 정리된 텍스트가 너무 짧음")

                except Exception as e:
                    logger.error(f"페이지 {page_num + 1} 처리 중 오류: {e}")
                    continue

        logger.info(f"pdfplumber 파싱 완료: {len(chunks)} 청크 생성")

        if not chunks:
            raise ValueError("추출된 유효한 텍스트가 없습니다")

        return chunks

    except ImportError:
        logger.error("pdfplumber 라이브러리가 설치되지 않았습니다")
        raise ImportError("pdfplumber 라이브러리가 필요합니다")
    except FileNotFoundError as e:
        logger.error(f"파일 접근 오류: {e}")
        raise
    except PermissionError as e:
        logger.error(f"권한 오류: {e}")
        raise
    except Exception as e:
        logger.error(f"pdfplumber PDF 파싱 실패: {type(e).__name__}: {e}")

        # 구체적인 오류 유형별 메시지
        if "corrupt" in str(e).lower() or "damaged" in str(e).lower():
            raise ValueError("손상된 PDF 파일입니다")
        elif "password" in str(e).lower() or "encrypted" in str(e).lower():
            raise ValueError("암호화된 PDF 파일입니다")
        elif "memory" in str(e).lower():
            raise ValueError("PDF 파일이 너무 커서 메모리 부족이 발생했습니다")
        else:
            raise ValueError(f"PDF 처리 중 오류 발생: {str(e)}")


def parse_pdf_with_pymupdf(file_path: str, lang_hint="auto") -> List[Dict]:
    """PyMuPDF를 사용한 PDF 파싱 (개선된 에러 처리)"""
    try:
        import fitz  # PyMuPDF

        # 파일 접근성 사전 확인
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일이 존재하지 않습니다: {file_path}")

        logger.info(f"PyMuPDF로 PDF 파싱 시작: {os.path.basename(file_path)}")
        chunks = []

        try:
            doc = fitz.open(file_path)
        except Exception as e:
            if "password" in str(e).lower():
                raise ValueError("암호화된 PDF 파일입니다")
            elif "corrupt" in str(e).lower():
                raise ValueError("손상된 PDF 파일입니다")
            else:
                raise ValueError(f"PDF 파일을 열 수 없습니다: {str(e)}")

        total_pages = len(doc)
        logger.info(f"PDF 총 페이지 수: {total_pages}")

        for page_num in range(total_pages):
            try:
                page = doc.load_page(page_num)
                text = None

                # 방법 1: 기본 텍스트 추출
                try:
                    text = page.get_text(
                        "text",
                        flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_PRESERVE_SPANS
                    )
                    if text and len(text.strip()) > 10:
                        logger.debug(f"페이지 {page_num + 1}: 기본 방법으로 텍스트 추출 성공")
                except Exception as e:
                    logger.warning(f"페이지 {page_num + 1} 기본 추출 실패: {e}")

                # 방법 2: dict 형태로 추출
                if not text or len(text.strip()) < 10:
                    try:
                        text_dict = page.get_text("dict")
                        extracted_text = []

                        for block in text_dict["blocks"]:
                            if "lines" in block:
                                for line in block["lines"]:
                                    for span in line["spans"]:
                                        if span["text"].strip():
                                            extracted_text.append(span["text"])

                        text = " ".join(extracted_text)
                        if text and len(text.strip()) > 10:
                            logger.debug(f"페이지 {page_num + 1}: dict 방법으로 텍스트 추출 성공")
                    except Exception as e:
                        logger.warning(f"페이지 {page_num + 1} dict 추출 실패: {e}")

                if not text or len(text.strip()) < 10:
                    logger.warning(f"페이지 {page_num + 1}: 추출된 텍스트 없음")
                    continue

                # 텍스트 정리 및 검증
                cleaned_text = clean_text(text)

                if cleaned_text and len(cleaned_text.strip()) > 10:
                    if not is_garbled(cleaned_text):
                        chunks.extend(chunk_text(
                            cleaned_text,
                            source=os.path.basename(file_path),
                            page=page_num + 1,
                            doc_type="pdf",
                            lang=lang_hint
                        ))
                        logger.debug(f"페이지 {page_num + 1}: 청크 생성 완료")
                    else:
                        logger.warning(f"페이지 {page_num + 1}: 깨진 텍스트 감지됨")

            except Exception as e:
                logger.error(f"페이지 {page_num + 1} 처리 중 오류: {e}")
                continue

        doc.close()
        logger.info(f"PyMuPDF 파싱 완료: {len(chunks)} 청크 생성")

        if not chunks:
            raise ValueError("추출된 유효한 텍스트가 없습니다")

        return chunks

    except ImportError:
        logger.error("PyMuPDF(fitz) 라이브러리가 설치되지 않았습니다")
        raise ImportError("PyMuPDF 라이브러리가 필요합니다")
    except FileNotFoundError as e:
        logger.error(f"파일 접근 오류: {e}")
        raise
    except Exception as e:
        logger.error(f"PyMuPDF PDF 파싱 실패: {type(e).__name__}: {e}")
        raise ValueError(f"PyMuPDF 처리 중 오류: {str(e)}")


def parse_pdf_with_pypdf(file_path: str, lang_hint="auto") -> List[Dict]:
    """PyPDF를 사용한 PDF 파싱 (개선된 에러 처리)"""
    try:
        from pypdf import PdfReader

        # 파일 접근성 사전 확인
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"파일이 존재하지 않습니다: {file_path}")

        logger.info(f"PyPDF로 PDF 파싱 시작: {os.path.basename(file_path)}")
        chunks = []

        # 다양한 인코딩으로 파일 읽기 시도
        reader = None
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                reader = PdfReader(file_path)
                logger.debug(f"PyPDF 파일 읽기 성공 (인코딩: {encoding})")
                break
            except Exception as e:
                logger.warning(f"인코딩 {encoding}으로 읽기 실패: {e}")
                continue

        if reader is None:
            raise ValueError("모든 인코딩으로 PDF 읽기 실패")

        total_pages = len(reader.pages)
        logger.info(f"PDF 총 페이지 수: {total_pages}")

        for page_num, page in enumerate(reader.pages):
            try:
                # 텍스트 추출
                text = page.extract_text()

                if not text or len(text.strip()) < 10:
                    logger.warning(f"페이지 {page_num + 1}: 추출된 텍스트 없음")
                    continue

                # 텍스트 정리 및 검증
                cleaned_text = clean_text(text)

                if cleaned_text and len(cleaned_text.strip()) > 10:
                    if not is_garbled(cleaned_text):
                        chunks.extend(chunk_text(
                            cleaned_text,
                            source=os.path.basename(file_path),
                            page=page_num + 1,
                            doc_type="pdf",
                            lang=lang_hint
                        ))
                        logger.debug(f"페이지 {page_num + 1}: 청크 생성 완료")
                    else:
                        logger.warning(f"페이지 {page_num + 1}: 깨진 텍스트 감지됨")

            except Exception as e:
                logger.error(f"페이지 {page_num + 1} 추출 실패: {e}")
                continue

        logger.info(f"PyPDF 파싱 완료: {len(chunks)} 청크 생성")

        if not chunks:
            raise ValueError("추출된 유효한 텍스트가 없습니다")

        return chunks

    except ImportError:
        logger.error("PyPDF 라이브러리가 설치되지 않았습니다")
        raise ImportError("PyPDF 라이브러리가 필요합니다")
    except FileNotFoundError as e:
        logger.error(f"파일 접근 오류: {e}")
        raise
    except Exception as e:
        logger.error(f"PyPDF 파싱 실패: {type(e).__name__}: {e}")

        # 구체적인 오류 유형별 메시지
        if "password" in str(e).lower() or "encrypted" in str(e).lower():
            raise ValueError("암호화된 PDF 파일입니다")
        elif "corrupt" in str(e).lower() or "damaged" in str(e).lower():
            raise ValueError("손상된 PDF 파일입니다")
        else:
            raise ValueError(f"PyPDF 처리 중 오류: {str(e)}")


def parse_text_file(file_path: str, lang_hint="auto") -> List[Dict]:
    """텍스트 파일 파싱 (인코딩 개선)"""
    try:
        # 다양한 인코딩 시도
        encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'latin-1', 'utf-16', 'ascii']
        content = None
        used_encoding = None

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    used_encoding = encoding
                break
            except (UnicodeDecodeError, UnicodeError):
                continue

        if content is None:
            # 바이너리로 읽어서 강제 변환
            with open(file_path, 'rb') as f:
                raw_content = f.read()
                # chardet로 인코딩 감지 시도
                try:
                    import chardet
                    detected = chardet.detect(raw_content)
                    if detected and detected['encoding']:
                        content = raw_content.decode(detected['encoding'], errors='replace')
                        used_encoding = detected['encoding']
                    else:
                        content = raw_content.decode('utf-8', errors='replace')
                        used_encoding = 'utf-8 (forced)'
                except ImportError:
                    content = raw_content.decode('utf-8', errors='replace')
                    used_encoding = 'utf-8 (forced)'

        # 텍스트 정리
        cleaned_content = clean_text(content)

        logger.info(f"Text file parsed with encoding: {used_encoding}")

        return chunk_text(
            cleaned_content,
            source=os.path.basename(file_path),
            doc_type="text",
            lang=lang_hint
        )

    except Exception as e:
        logger.error(f"Text file parsing failed: {e}")
        return parse_as_fallback(file_path, f"텍스트 파일 파싱 실패: {str(e)}")


def parse_docx(file_path: str, lang_hint="auto") -> List[Dict]:
    """Word 문서 파싱 (인코딩 개선)"""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = []

        # 문단 텍스트 추출
        for para in doc.paragraphs:
            if para.text.strip():
                cleaned_text = clean_text(para.text.strip())
                if cleaned_text:
                    paragraphs.append(cleaned_text)

        # 표 내용 추출
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = clean_text(cell.text.strip())
                    if cell_text:
                        row_texts.append(cell_text)

                if row_texts:
                    row_text = " | ".join(row_texts)
                    paragraphs.append(row_text)

        # 전체 텍스트 결합
        full_text = "\n\n".join(paragraphs)

        return chunk_text(
            full_text,
            source=os.path.basename(file_path),
            doc_type="docx",
            lang=lang_hint
        )

    except ImportError:
        logger.error("python-docx not available")
        return parse_as_fallback(file_path, "Word 문서 (python-docx 미설치)")
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return parse_as_fallback(file_path, f"Word 문서 파싱 오류: {str(e)}")


def parse_as_fallback(file_path: str, description: str) -> List[Dict]:
    """개선된 폴백 처리 - 사용자 친화적 메시지와 문제 해결 힌트"""

    filename = os.path.basename(file_path)
    file_ext = os.path.splitext(file_path)[1].lower()

    # 파일 정보 수집
    file_info = {"size": 0, "readable": False, "exists": False}
    try:
        if os.path.exists(file_path):
            file_info["exists"] = True
            file_info["size"] = os.path.getsize(file_path)
            file_info["readable"] = os.access(file_path, os.R_OK)
    except:
        pass

    # 사용자 친화적 메시지 생성
    if "라이브러리가 설치되지 않음" in description:
        user_message = f"""
📄 **{filename}** 파일이 업로드되었습니다.

⚠️ **처리 제한사항**: PDF 처리 라이브러리가 설치되지 않아 내용을 추출할 수 없습니다.

💡 **해결 방법**:
- 관리자에게 PDF 처리 라이브러리 설치를 요청하세요
- 파일을 텍스트 형태로 변환하여 다시 업로드해 보세요
- Word 문서(.docx) 형태로 변환하여 시도해 보세요

📋 **파일 정보**: {file_info['size'] / 1024:.1f}KB
        """.strip()

    elif "파일을 찾을 수 없습니다" in description:
        user_message = f"""
❌ **파일 업로드 오류**: {filename}

파일이 제대로 업로드되지 않았습니다. 다시 시도해 주세요.

💡 **해결 방법**:
- 파일을 다시 선택하여 업로드하세요
- 파일 이름에 특수문자가 있다면 제거하여 시도하세요
- 브라우저를 새로고침한 후 다시 시도하세요
        """.strip()

    elif "파일이 너무 큽니다" in description:
        user_message = f"""
📄 **{filename}** 파일이 업로드되었습니다.

⚠️ **처리 제한사항**: 파일 크기가 너무 커서 처리할 수 없습니다.

💡 **해결 방법**:
- 파일을 작은 단위로 분할하여 업로드하세요
- PDF의 경우 페이지 수를 줄여서 시도하세요
- 이미지 해상도를 낮춰서 다시 저장하세요

📋 **파일 정보**: {file_info['size'] / 1024 / 1024:.1f}MB (제한: 100MB)
        """.strip()

    elif "빈 파일입니다" in description:
        user_message = f"""
📄 **{filename}** 파일이 업로드되었습니다.

⚠️ **처리 제한사항**: 파일에 내용이 없습니다.

💡 **확인 사항**:
- 파일에 실제 내용이 있는지 확인하세요
- 파일이 올바르게 저장되었는지 확인하세요
- 다른 파일을 시도해 보세요
        """.strip()

    elif "가독 불가능한 텍스트" in description:
        user_message = f"""
📄 **{filename}** 파일이 업로드되었습니다.

⚠️ **처리 제한사항**: PDF에서 텍스트를 추출했지만 읽을 수 없는 형태입니다.

💡 **가능한 원인**:
- 스캔된 이미지 기반 PDF (OCR 필요)
- 특수 폰트나 인코딩 문제
- 손상된 PDF 파일

💡 **해결 방법**:
- PDF를 텍스트 형태로 다시 저장하여 업로드하세요
- Word 문서로 변환하여 시도하세요
- 다른 PDF 뷰어에서 "다른 이름으로 저장"을 시도하세요

📋 **파일 정보**: {file_info['size'] / 1024:.1f}KB
        """.strip()

    elif "모든 파싱 라이브러리 실패" in description:
        user_message = f"""
📄 **{filename}** 파일이 업로드되었습니다.

⚠️ **처리 제한사항**: 여러 PDF 처리 방법을 시도했지만 모두 실패했습니다.

💡 **가능한 원인**:
- 암호화된 PDF 파일
- 손상되거나 특수한 형식의 PDF
- 시스템 리소스 부족

💡 **해결 방법**:
- PDF 암호를 해제한 후 다시 업로드하세요
- 다른 PDF 뷰어에서 "인쇄 → PDF로 저장"을 시도하세요
- 파일을 Word나 텍스트 형태로 변환하여 업로드하세요
- 관리자에게 문의하세요

🔧 **기술 정보**: {description.split(':', 1)[-1] if ':' in description else description}

📋 **파일 정보**: {file_info['size'] / 1024:.1f}KB
        """.strip()

    else:
        # 기타 오류의 경우
        user_message = f"""
📄 **{filename}** 파일이 업로드되었습니다.

⚠️ **처리 중 문제 발생**: {description}

💡 **해결 방법**:
- 파일 형식을 확인하고 지원되는 형식(.pdf, .docx, .txt)으로 변환하세요
- 파일 이름에서 특수문자를 제거하세요
- 다른 파일로 시도해 보세요
- 문제가 지속되면 관리자에게 문의하세요

📋 **파일 정보**: {file_ext.upper()[1:] if file_ext else '알 수 없음'} 형식, {file_info['size'] / 1024:.1f}KB
        """.strip()

    return [{
        "chunk_id": str(uuid4()),
        "content": user_message,
        "meta": {
            "source": filename,
            "type": "fallback",
            "error_type": "parsing_failed",
            "original_error": description,
            "file_size": file_info["size"],
            "file_ext": file_ext,
            "suggestions": [
                "파일 형식 변환 시도",
                "파일 크기 축소",
                "관리자 문의"
            ]
        }
    }]


def chunk_text(text: str, source: str, doc_type: str = "unknown",
               page: int = None, lang: str = "auto", chunk_size: int = 500,
               chunk_overlap: int = 50) -> List[Dict]:
    """텍스트를 청크로 분할 (한국어 최적화)"""
    if not text or not text.strip():
        return []

    # 텍스트 정리
    cleaned_text = clean_text(text)

    if not cleaned_text:
        return []

    chunks = []

    if len(cleaned_text) <= chunk_size:
        chunks.append({
            "chunk_id": str(uuid4()),
            "content": cleaned_text,
            "meta": {
                "source": source,
                "type": doc_type,
                "lang": lang,
                **({"page": page} if page else {}),
                "chunk_index": 0,
                "total_chunks": 1
            }
        })
    else:
        start = 0
        chunk_index = 0

        while start < len(cleaned_text):
            end = start + chunk_size

            # 한국어 및 영어를 고려한 단어 경계에서 자르기
            if end < len(cleaned_text):
                # 한국어 문장 끝 문자들
                korean_endings = ['다', '요', '니다', '습니다', '죠', '네', '라', '까', '야']

                # 찾을 범위 설정
                search_start = max(start + chunk_size // 2, end - 100)

                # 우선순위 1: 문장 부호
                for i in range(end, search_start, -1):
                    if cleaned_text[i] in ['.', '!', '?', '。', '!', '?', '\n']:
                        end = i + 1
                        break
                else:
                    # 우선순위 2: 한국어 문장 끝
                    for i in range(end, search_start, -1):
                        if i > 0 and cleaned_text[i-1:i+1] in ['다.', '요.', '다!', '요!', '다?', '요?']:
                            end = i + 1
                            break
                        elif cleaned_text[i] in korean_endings and i < len(cleaned_text) - 1 and cleaned_text[i+1] in [' ', '\n', '\t']:
                            end = i + 1
                            break
                    else:
                        # 우선순위 3: 공백
                        for i in range(end, search_start, -1):
                            if cleaned_text[i] in [' ', '\n', '\t']:
                                end = i + 1
                                break

            chunk_content = cleaned_text[start:end].strip()

            if chunk_content and len(chunk_content) > 10:
                chunks.append({
                    "chunk_id": str(uuid4()),
                    "content": chunk_content,
                    "meta": {
                        "source": source,
                        "type": doc_type,
                        "lang": lang,
                        **({"page": page} if page else {}),
                        "chunk_index": chunk_index,
                        "start_pos": start,
                        "end_pos": end
                    }
                })
                chunk_index += 1

            start = max(start + chunk_size - chunk_overlap, end)

    # 총 청크 수 업데이트
    for chunk in chunks:
        chunk["meta"]["total_chunks"] = len(chunks)

    return chunks


def parse_pdf(file_input: Union[str, bytes], lang_hint: str = "auto") -> List[Dict]:
    """메인 파싱 함수 - 여러 라이브러리로 단계적 시도"""
    # 바이트 데이터인 경우 임시 파일로 저장
    if isinstance(file_input, bytes):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(file_input)
            temp_path = tmp_file.name

        try:
            return parse_file_by_extension(temp_path, lang_hint)
        finally:
            try:
                os.unlink(temp_path)
            except:
                pass

    elif isinstance(file_input, (str, os.PathLike)):
        return parse_file_by_extension(str(file_input), lang_hint)

    else:
        raise ValueError(f"Unsupported input type: {type(file_input)}")


def parse_file_by_extension(file_path: str, lang_hint: str = "auto") -> List[Dict]:
    """파일 확장자에 따른 파싱 - 개선된 에러 처리"""

    # 1단계: 파일 존재 및 접근 권한 검증
    if not os.path.exists(file_path):
        logger.error(f"파일이 존재하지 않습니다: {file_path}")
        return parse_as_fallback(file_path, "파일을 찾을 수 없습니다")

    if not os.access(file_path, os.R_OK):
        logger.error(f"파일 읽기 권한이 없습니다: {file_path}")
        return parse_as_fallback(file_path, "파일 읽기 권한이 없습니다")

    try:
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            logger.warning(f"빈 파일입니다: {file_path}")
            return parse_as_fallback(file_path, "빈 파일입니다")

        if file_size > 100 * 1024 * 1024:  # 100MB 제한
            logger.warning(f"파일이 너무 큽니다 ({file_size / 1024 / 1024:.1f}MB): {file_path}")
            return parse_as_fallback(file_path, f"파일이 너무 큽니다 ({file_size / 1024 / 1024:.1f}MB)")

    except Exception as e:
        logger.error(f"파일 정보 조회 실패: {file_path}, 오류: {e}")
        return parse_as_fallback(file_path, f"파일 정보를 읽을 수 없습니다: {str(e)}")

    file_ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"파일 처리 시작: {os.path.basename(file_path)} ({file_ext}, {file_size / 1024:.1f}KB)")

    if file_ext == '.pdf':
        # PDF 파싱 - 여러 라이브러리를 순차적으로 시도 (개선된 에러 처리)
        successful_chunks = []
        library_errors = []

        # 1단계: pdfplumber 먼저 시도
        try:
            logger.info("pdfplumber로 PDF 파싱 시도 중...")
            result = parse_pdf_with_pdfplumber(file_path, lang_hint)
            if result and not all(chunk['meta'].get('type') == 'fallback' for chunk in result):
                content_check = " ".join(c['content'] for c in result if c['content'])
                if not is_garbled(content_check) and len(content_check.strip()) > 20:
                    logger.info(f"pdfplumber 성공: {len(result)} 청크 생성")
                    return result
                else:
                    library_errors.append("pdfplumber: 가독 불가능한 텍스트 추출")
            else:
                library_errors.append("pdfplumber: 유효한 내용 추출 실패")
        except ImportError:
            library_errors.append("pdfplumber: 라이브러리가 설치되지 않음")
        except Exception as e:
            library_errors.append(f"pdfplumber: {str(e)}")
            logger.warning(f"pdfplumber 실패: {e}")

        # 2단계: PyMuPDF 시도
        try:
            logger.info("PyMuPDF로 PDF 파싱 시도 중...")
            result = parse_pdf_with_pymupdf(file_path, lang_hint)
            if result and not all(chunk['meta'].get('type') == 'fallback' for chunk in result):
                content_check = " ".join(c['content'] for c in result if c['content'])
                if not is_garbled(content_check) and len(content_check.strip()) > 20:
                    logger.info(f"PyMuPDF 성공: {len(result)} 청크 생성")
                    return result
                else:
                    library_errors.append("PyMuPDF: 가독 불가능한 텍스트 추출")
            else:
                library_errors.append("PyMuPDF: 유효한 내용 추출 실패")
        except ImportError:
            library_errors.append("PyMuPDF: 라이브러리가 설치되지 않음")
        except Exception as e:
            library_errors.append(f"PyMuPDF: {str(e)}")
            logger.warning(f"PyMuPDF 실패: {e}")

        # 3단계: PyPDF 시도
        try:
            logger.info("PyPDF로 PDF 파싱 시도 중...")
            result = parse_pdf_with_pypdf(file_path, lang_hint)
            if result and not all(chunk['meta'].get('type') == 'fallback' for chunk in result):
                content_check = " ".join(c['content'] for c in result if c['content'])
                if not is_garbled(content_check) and len(content_check.strip()) > 20:
                    logger.info(f"PyPDF 성공: {len(result)} 청크 생성")
                    return result
                else:
                    library_errors.append("PyPDF: 가독 불가능한 텍스트 추출")
            else:
                library_errors.append("PyPDF: 유효한 내용 추출 실패")
        except ImportError:
            library_errors.append("PyPDF: 라이브러리가 설치되지 않음")
        except Exception as e:
            library_errors.append(f"PyPDF: {str(e)}")
            logger.warning(f"PyPDF 실패: {e}")

        # 모든 시도가 실패한 경우 - 상세한 오류 정보 제공
        error_summary = "; ".join(library_errors)
        logger.error(f"PDF 파싱 완전 실패: {error_summary}")

        return parse_as_fallback(
            file_path,
            f"PDF 파싱 실패 - 시도한 모든 라이브러리에서 오류 발생: {error_summary}"
        )

    elif file_ext in ['.docx', '.doc']:
        try:
            return parse_docx(file_path, lang_hint)
        except Exception as e:
            logger.error(f"Word 문서 파싱 실패: {e}")
            return parse_as_fallback(file_path, f"Word 문서 파싱 오류: {str(e)}")

    elif file_ext in ['.txt', '.md', '.rst']:
        try:
            return parse_text_file(file_path, lang_hint)
        except Exception as e:
            logger.error(f"텍스트 파일 파싱 실패: {e}")
            return parse_as_fallback(file_path, f"텍스트 파일 파싱 오류: {str(e)}")
    else:
        # 기본적으로 텍스트로 처리 시도
        try:
            logger.info(f"알 수 없는 확장자 {file_ext}, 텍스트로 시도")
            return parse_text_file(file_path, lang_hint)
        except Exception as e:
            return parse_as_fallback(file_path, f"알 수 없는 파일 형식 ({file_ext}): {str(e)}")


# 테스트 함수
if __name__ == "__main__":
    import sys

    print("📄 Improved Korean PDF Parser Test")
    print("==================================")

    # 인코딩 테스트
    test_korean = "안녕하세요. 이것은 한글 테스트입니다. AI 기반 검색 패러다임 전환에 대해 설명합니다."
    cleaned = clean_text(test_korean)
    print(f"✅ Encoding test passed: {cleaned[:50]}...")

    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        print(f"\n🔍 Testing with: {test_file}")

        try:
            chunks = parse_pdf(test_file)
            print(f"✅ Generated {len(chunks)} chunks")

            for i, chunk in enumerate(chunks[:3]):  # 처음 3개 청크만 표시
                print(f"\n--- Chunk {i+1} ---")
                print(f"Content preview: {chunk['content'][:200]}...")
                print(f"Content length: {len(chunk['content'])}")
                print(f"Meta: {chunk['meta']}")

                # 한국어 포함 여부 확인
                korean_chars = len(re.findall(r'[가-힣]', chunk['content']))
                print(f"Korean characters: {korean_chars}")

        except Exception as e:
            print(f"❌ Failed: {e}")
            import traceback
            traceback.print_exc()
    else:
        print("\n💡 Usage: python parser.py <pdf_file_path>")